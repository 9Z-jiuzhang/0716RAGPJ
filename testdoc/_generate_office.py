"""Generate DOCX/PDF sample files into /out (or current directory)."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUT = Path("/out") if Path("/out").exists() else Path(__file__).resolve().parent


def write_docx() -> Path:
    doc = Document()
    doc.add_heading("信息安全与账号权限规范", level=1)
    p = doc.add_paragraph("文档类型：Word DOCX（.docx）\n用于验证 Word 解析、分段与向量化流水线。")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("1. 账号开通", level=2)
    doc.add_paragraph("新员工入职后由 IT 开通企业邮箱、OA、VPN 与代码仓库权限。")
    doc.add_paragraph("临时外包账号有效期默认 90 天，到期自动冻结。")

    doc.add_heading("2. 权限申请", level=2)
    for item in (
        "生产环境数据库只读权限须由部门负责人与 DBA 双审批。",
        "禁止共享个人账号；发现共享立即停用并纳入安全审计。",
        "离职当日回收全部系统权限，硬件资产在 3 个工作日内归还。",
    ):
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("3. 密码策略", level=2)
    doc.add_paragraph("密码长度不少于 12 位，需包含大小写字母、数字与特殊字符；每 90 天强制轮换。")
    doc.add_paragraph("关键词：账号权限、VPN、双审批、安全审计、密码策略")

    path = OUT / "03_信息安全与账号权限规范.docx"
    doc.save(path)
    return path


def write_pdf() -> Path:
    """手写一份最小可用 PDF（含中英文），不依赖额外字体库。"""
    # 使用 PDFDocEncoding 可表示的文本：英文为主，并附带可检索的 ASCII 关键词；
    # 中文完整内容见同目录 TXT/MD/DOCX，本 PDF 重点验证解析链路。
    lines = [
        "IT Account & Security Policy (PDF Test Document)",
        "",
        "File type: PDF (.pdf)",
        "Purpose: validate PDF text extraction, chunking and embedding.",
        "",
        "1. Account Provisioning",
        "- New hires receive Email, OA, VPN and repo access from IT.",
        "- Contractor accounts expire in 90 days by default.",
        "",
        "2. Access Control",
        "- Production DB read-only access needs manager + DBA dual approval.",
        "- Sharing personal accounts is forbidden and triggers audit freeze.",
        "- All access must be revoked on the employee's last working day.",
        "",
        "3. Password Policy",
        "- Minimum 12 characters with upper/lower/digit/special.",
        "- Force rotation every 90 days.",
        "",
        "Keywords: VPN, OA, dual-approval, password-policy, security-audit",
        "Related Chinese docs: 信息安全与账号权限规范 (DOCX) / 年假与加班制度 (TXT)",
    ]

    # Build a simple single-page PDF with Helvetica
    content_lines = ["BT", "/F1 11 Tf", "50 780 Td", "14 TL"]
    first = True
    for line in lines:
        safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        if first:
            content_lines.append(f"({safe}) Tj")
            first = False
        else:
            content_lines.append("T*")
            content_lines.append(f"({safe}) Tj")
    content_lines.append("ET")
    stream = "\n".join(content_lines).encode("latin-1", errors="replace")

    objects: list[bytes] = []
    objects.append(b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n")
    objects.append(b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n")
    objects.append(
        b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>endobj\n"
    )
    objects.append(
        f"4 0 obj<< /Length {len(stream)} >>stream\n".encode("ascii")
        + stream
        + b"\nendstream\nendobj\n"
    )
    objects.append(b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n")

    buf = io.BytesIO()
    buf.write(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(buf.tell())
        buf.write(obj)
    xref_pos = buf.tell()
    buf.write(f"xref\n0 {len(offsets)}\n".encode("ascii"))
    buf.write(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        buf.write(f"{off:010d} 00000 n \n".encode("ascii"))
    buf.write(
        f"trailer<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode(
            "ascii"
        )
    )

    path = OUT / "04_信息安全策略_PDF测试.pdf"
    path.write_bytes(buf.getvalue())
    return path


def write_doc_as_docx_alias() -> Path:
    """额外提供一份扩展名为 .doc、内容实为 docx 的样例，覆盖误标解析路径。"""
    doc = Document()
    doc.add_heading("员工入职与试用期手册（DOC 兼容样例）", level=1)
    doc.add_paragraph("说明：本文件扩展名为 .doc，内部为 OOXML（docx）结构，用于验证解析器兼容分支。")
    doc.add_heading("试用期", level=2)
    doc.add_paragraph("试用期一般为 3 个月，期满前 5 个工作日完成转正评估。")
    doc.add_heading("入职材料", level=2)
    doc.add_paragraph("需提交身份证、学历证明、银行卡与体检报告；材料不全不得开通生产系统权限。")
    doc.add_paragraph("关键词：入职、试用期、转正评估、材料提交")
    path = OUT / "06_入职试用期手册_docx结构.doc"
    doc.save(path)
    return path


if __name__ == "__main__":
    paths = [write_docx(), write_pdf(), write_doc_as_docx_alias()]
    for p in paths:
        print("wrote", p.name, p.stat().st_size)
